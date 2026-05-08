import os
import io
import zipfile
import base64
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS

app = Flask(__name__)
CORS(app)

class DocxOLEEmbedder:
    def __init__(self):
        self.temp_dir = os.path.join(os.getcwd(), 'temp')
        os.makedirs(self.temp_dir, exist_ok=True)
        self.use_win32 = False
        self.word = None
        self.excel = None
        
    def init_com(self):
        try:
            import pythoncom
            from win32com.client import Dispatch
            
            pythoncom.CoInitialize()
            
            print("[COM] 正在初始化Word/WPS...")
            try:
                self.word = Dispatch('WPS.Application')
                self.word.Visible = False
                print("[COM] ✅ WPS.Application 初始化成功")
            except Exception as e:
                try:
                    self.word = Dispatch('Kwps.Application')
                    self.word.Visible = False
                    print("[COM] ✅ Kwps.Application 初始化成功")
                except Exception as e2:
                    try:
                        self.word = Dispatch('Word.Application')
                        self.word.Visible = False
                        print("[COM] ✅ Word.Application 初始化成功")
                    except Exception as e3:
                        print(f"[COM] ❌ 所有Word/WPS初始化失败: {e3}")
                        return False
            
            print("[COM] 正在初始化Excel/WPS Excel...")
            try:
                self.excel = Dispatch('Kexcel.Application')
                self.excel.Visible = False
                print("[COM] ✅ Kexcel.Application 初始化成功")
            except Exception as e:
                try:
                    self.excel = Dispatch('Excel.Application')
                    self.excel.Visible = False
                    print("[COM] ✅ Excel.Application 初始化成功")
                except Exception as e2:
                    print(f"[COM] ⚠️ Excel初始化失败（不是必需的）: {e2}")
            
            self.use_win32 = True
            print("[COM] ✅ COM模式初始化完成")
            return True
        except ImportError:
            print("pywin32 not installed, using pure Python mode")
            return False
    
    def cleanup_com(self):
        if self.word:
            try:
                self.word.Quit()
            except:
                pass
            self.word = None
        if self.excel:
            try:
                self.excel.Quit()
            except:
                pass
            self.excel = None
        try:
            import pythoncom
            pythoncom.CoUninitialize()
        except:
            pass
    
    def embed_excel_to_docx(self, docx_bytes, excel_files, placeholder_mappings):
        if not self.use_win32:
            return self.embed_excel_pure_python(docx_bytes, excel_files, placeholder_mappings)
        
        temp_docx_path = os.path.join(self.temp_dir, 'template.docx')
        with open(temp_docx_path, 'wb') as f:
            f.write(docx_bytes)
        
        excel_file_list = list(excel_files.items())
        excel_index = 0
        
        try:
            doc = self.word.Documents.Open(temp_docx_path)
            print("[OLE] 文档已打开")
            
            placeholders_found = []
            try:
                find_range = doc.Content.Find
                find_range.Text = "{{excel:"
                find_range.Wrap = 1
                
                while find_range.Execute():
                    full_text = find_range.Parent.Text
                    end_pos = full_text.find("}}", find_range.Parent.Start)
                    if end_pos > 0:
                        placeholder = full_text[:end_pos + 2]
                        placeholders_found.append(placeholder)
                    
                    find_range.Parent.Start = find_range.Parent.End + 1
            except Exception as e:
                print(f"[OLE] 查找占位符失败: {e}")
            
            print(f"[OLE] 找到 {len(placeholders_found)} 个占位符")
            
            for i, placeholder in enumerate(placeholders_found):
                if i >= len(excel_file_list):
                    print("[OLE] ⚠️ Excel文件不足")
                    break
                
                excel_filename, excel_base64 = excel_file_list[i]
                print(f"[OLE] 处理占位符: {placeholder} -> {excel_filename}")
                
                try:
                    find_range = doc.Content.Find
                    find_range.Text = placeholder
                    find_range.Wrap = 1
                    
                    if find_range.Execute():
                        selection = self.word.Selection
                        selection.Range = find_range.Parent
                        selection.Delete()
                        
                        temp_excel_path = os.path.join(self.temp_dir, excel_filename)
                        excel_bytes_data = base64.b64decode(excel_base64)
                        with open(temp_excel_path, 'wb') as f:
                            f.write(excel_bytes_data)
                        
                        try:
                            ole_obj = doc.InlineShapes.AddOLEObject(
                                ClassType="",
                                FileName=temp_excel_path,
                                LinkToFile=False,
                                DisplayAsIcon=True
                            )
                            print(f"[OLE] ✅ OLE对象插入成功")
                            os.remove(temp_excel_path)
                        except Exception as e:
                            print(f"[OLE] ❌ OLE插入失败: {e}")
                            os.remove(temp_excel_path)
                except Exception as e:
                    print(f"[OLE] 处理失败: {e}")
            
            print(f"[OLE] 共处理 {excel_index} 个Excel文件")
            
            output_path = os.path.join(self.temp_dir, 'output.docx')
            doc.SaveAs(output_path)
            doc.Close()
            
            with open(output_path, 'rb') as f:
                result_bytes = f.read()
            
            os.remove(output_path)
            return result_bytes
            
        except Exception as e:
            print(f"Error in embed_excel_to_docx: {e}")
            return docx_bytes
    
    def embed_excel_com(self, docx_bytes, excel_files, placeholder_mappings):
        print("[OLE] COM方式已禁用，直接返回原文档")
        return docx_bytes
    
    def embed_excel_libreoffice(self, docx_bytes, excel_files):
        try:
            import os
            import tempfile
            import subprocess
            
            with tempfile.TemporaryDirectory() as temp_dir:
                temp_doc_path = os.path.join(temp_dir, 'template.docx')
                with open(temp_doc_path, 'wb') as f:
                    f.write(docx_bytes)
                
                excel_info = []
                for excel_filename, excel_base64 in excel_files.items():
                    excel_bytes_data = base64.b64decode(excel_base64)
                    temp_excel_path = os.path.join(temp_dir, excel_filename)
                    with open(temp_excel_path, 'wb') as f:
                        f.write(excel_bytes_data)
                    
                    identifier = excel_filename.split('.')[0]
                    excel_info.append({
                        'filename': excel_filename,
                        'path': temp_excel_path,
                        'identifier': identifier
                    })
                
                output_doc_path = os.path.join(temp_dir, 'output.docx')
                
                ps_script = self.generate_powershell_script(temp_doc_path, output_doc_path, excel_info)
                ps_path = os.path.join(temp_dir, 'embed_ole.ps1')
                log_path = os.path.join(temp_dir, 'ps_output.log')
                with open(ps_path, 'w', encoding='utf-8-sig') as f:
                    f.write(ps_script)
                
                print(f"[DEBUG] PowerShell script saved to: {ps_path}")
                print(f"[DEBUG] Script content preview:\n{ps_script[:500]}...")
                
                cmd = f'powershell -ExecutionPolicy Bypass -File "{ps_path}"'
                
                print(f"[DEBUG] Executing: {cmd}")
                
                try:
                    result = subprocess.run(cmd, shell=True, capture_output=True, text=True, encoding='utf-8', timeout=120)
                    ps_output = result.stdout + result.stderr
                    print(f"[DEBUG] PowerShell exit code: {result.returncode}")
                    print(f"[OLE] PowerShell输出:\n{ps_output}")
                    
                    if os.path.exists(log_path):
                        with open(log_path, 'r', encoding='utf-8-sig', errors='replace') as f:
                            log_output = f.read()
                        print(f"[OLE] 日志文件内容:\n{log_output}")
                    else:
                        print(f"[DEBUG] 日志文件不存在: {log_path}")
                except subprocess.TimeoutExpired:
                    print("[OLE] PowerShell超时")
                    return docx_bytes
                
                if os.path.exists(output_doc_path):
                    with open(output_doc_path, 'rb') as f:
                        result_bytes = f.read()
                    print(f"[OLE] PowerShell COM方式成功")
                    return result_bytes
                else:
                    print("[OLE] PowerShell未生成输出文件")
                    return docx_bytes
            
        except Exception as e:
            print(f"PowerShell COM embedding failed: {e}")
            return docx_bytes
    
    def generate_powershell_script(self, input_path, output_path, excel_info):
        script_lines = [
            '# 设置控制台编码为UTF-8',
            '[Console]::OutputEncoding = [System.Text.Encoding]::UTF8',
            '[Console]::InputEncoding = [System.Text.Encoding]::UTF8',
            '$OutputEncoding = [System.Text.Encoding]::UTF8',
            '',
            'Write-Host "=== 开始处理 ==="',
            '',
            '$word = $null',
            '$successCount = 0',
            '$totalCount = ' + str(len(excel_info)),
            '',
            'try {',
            '    Write-Host "尝试启动WPS.Application..."',
            '    $word = New-Object -ComObject WPS.Application',
            '    Write-Host "成功启动WPS.Application"',
            '} catch {',
            '    Write-Host "WPS.Application失败: $_"',
            '    try {',
            '        Write-Host "尝试启动Word.Application..."',
            '        $word = New-Object -ComObject Word.Application',
            '        Write-Host "成功启动Word.Application"',
            '    } catch {',
            '        Write-Host "Word.Application失败: $_"',
            '        Write-Host "错误: 无法启动WPS或Word"',
            '        exit 1',
            '    }',
            '}',
            '',
            '$word.Visible = $false',
            '$word.DisplayAlerts = 0',
            '',
            'Write-Host "输入文档: ' + input_path.replace('\\', '/') + '"',
            '',
            'try {',
            '    $doc = $word.Documents.Open("' + input_path.replace('\\', '\\\\') + '")',
            '    Write-Host "成功打开文档"',
            '} catch {',
            '    Write-Host "打开文档失败: $_"',
            '    $word.Quit()',
            '    exit 1',
            '}',
            '',
            'Write-Host "Excel文件数量: ' + str(len(excel_info)) + '"',
            ''
        ]
        
        for idx, info in enumerate(excel_info):
            placeholder = f'{{{{excel:{info["identifier"]}}}}}'
            excel_path_escaped = info['path'].replace('\\', '\\\\')
            filename = info['filename']
            
            script_lines.append('')
            script_lines.append(f'Write-Host "--- 处理第{idx+1}个Excel文件: {filename} ---"')
            script_lines.append(f'$findText = "{placeholder}"')
            script_lines.append(f'$excelPath = "{excel_path_escaped}"')
            script_lines.append('Write-Host "占位符: $findText"')
            script_lines.append('Write-Host "Excel路径: $excelPath"')
            script_lines.append('')
            script_lines.append('if (-not (Test-Path $excelPath)) {')
            script_lines.append('    Write-Host "错误: Excel文件不存在"')
            script_lines.append('    continue')
            script_lines.append('}')
            script_lines.append('')
            script_lines.append('Write-Host "步骤A: 开始处理占位符..."')
            script_lines.append('$count = 0')
            script_lines.append('Write-Host "步骤B: 将光标移到文档开头"')
            script_lines.append('try {')
            script_lines.append('    $word.Selection.HomeKey(6)')
            script_lines.append('    Write-Host "成功移动光标"')
            script_lines.append('} catch {')
            script_lines.append('    Write-Host "移动光标失败: $_"')
            script_lines.append('}')
            script_lines.append('$found = $true')
            script_lines.append('$loopCount = 0')
            script_lines.append('Write-Host "步骤C: 开始循环..."')
            script_lines.append('while ($found -and $loopCount -lt 100) {')
            script_lines.append('    $loopCount++')
            script_lines.append('    Write-Host "循环次数: $loopCount"')
            script_lines.append('    $word.Selection.Find.ClearFormatting()')
            script_lines.append('    $word.Selection.Find.Text = $findText')
            script_lines.append('    $word.Selection.Find.Forward = $true')
            script_lines.append('    $word.Selection.Find.Wrap = 2')
            script_lines.append('    $word.Selection.Find.MatchCase = $false')
            script_lines.append('    $word.Selection.Find.MatchWholeWord = $false')
            script_lines.append('    $word.Selection.Find.MatchWildcards = $false')
            script_lines.append('')
            script_lines.append('    $found = $word.Selection.Find.Execute()')
            script_lines.append('')
            script_lines.append('    if ($found) {')
            script_lines.append('        Write-Host "找到占位符!"')
            script_lines.append('        try {')
            script_lines.append('            Write-Host "插入OLE对象..."')
            script_lines.append('            # 获取当前选择范围')
            script_lines.append('            $selection = $word.Selection')
            script_lines.append('            # 获取范围对象')
            script_lines.append('            $range = $selection.Range')
            script_lines.append('            # 使用Shapes.AddOLEObject并设置图标显示')
            script_lines.append('            $shape = $range.InlineShapes.AddOLEObject("Excel.Sheet.12", $excelPath, $false, $true, "", 0, "Excel")')
            script_lines.append('            # 确保显示为图标')
            script_lines.append('            if ($shape.OLEFormat) {')
            script_lines.append('                $shape.OLEFormat.DisplayAsIcon = $true')
            script_lines.append('            }')
            script_lines.append('            Write-Host "成功插入OLE对象"')
            script_lines.append('            $count++')
            script_lines.append('            $successCount++')
            script_lines.append('        } catch {')
            script_lines.append('            Write-Host "插入OLE失败: $_"')
            script_lines.append('            Write-Host "错误详情: $($_.Exception.Message)"')
            script_lines.append('        }')
            script_lines.append('    } else {')
            script_lines.append('        Write-Host "未找到占位符，结束循环"')
            script_lines.append('    }')
            script_lines.append('}')
            script_lines.append('')
            script_lines.append('Write-Host "共处理 $count 个占位符"')
            script_lines.append('')
        
        script_lines.extend([
            'Write-Host ""',
            'Write-Host "=== 调试信息 ==="',
            'Write-Host "文档内容预览(前2000字符):"',
            'Write-Host $doc.Content.Text.Substring(0, [Math]::Min(2000, $doc.Content.Text.Length))',
            '',
            'Write-Host "=== 保存文档 ==="',
            'Write-Host "输出文档: ' + output_path.replace('\\', '/') + '"',
            '',
            'try {',
            '    $doc.SaveAs("' + output_path.replace('\\', '\\\\') + '")',
            '    Write-Host "成功保存文档"',
            '} catch {',
            '    Write-Host "保存文档失败: $_"',
            '}',
            '',
            '$doc.Close()',
            'Write-Host "关闭文档"',
            '',
            '$word.Quit()',
            '[System.Runtime.Interopservices.Marshal]::ReleaseComObject($word) | Out-Null',
            'Write-Host "退出Word"',
            '',
            'Write-Host "=== 完成 ==="',
            'Write-Host "成功: $successCount / $totalCount"'
        ])
        
        full_script = '\n'.join(script_lines)
        print("[DEBUG] Generated PowerShell script:")
        print(full_script[:1000] + "..." if len(full_script) > 1000 else full_script)
        
        return full_script
    
    def embed_excel_pure_python(self, docx_bytes, excel_files, placeholder_mappings):
        try:
            import zipfile
            import re
            import olefile
            
            template_zip = zipfile.ZipFile(io.BytesIO(docx_bytes))
            
            document_xml = ''
            rels_content = ''
            max_rid = 1
            
            for name in template_zip.namelist():
                content = template_zip.read(name)
                
                if name == 'word/document.xml':
                    document_xml = content.decode('utf-8')
                elif name == 'word/_rels/document.xml.rels':
                    rels_content = content.decode('utf-8')
                    rid_match = re.findall(r'rId(\d+)', rels_content)
                    if rid_match:
                        max_rid = max([int(r) for r in rid_match]) + 1
            
            excel_file_list = list(excel_files.items())
            
            if not excel_file_list:
                print("[OLE] 没有Excel文件")
                return docx_bytes
            
            excel_placeholders = re.findall(r'\{\{excel:([^}]+)\}\}', document_xml)
            print(f"[OLE] 找到 {len(excel_placeholders)} 个占位符")
            
            if not excel_placeholders:
                print("[OLE] 没有找到占位符")
                return docx_bytes
            
            output_buffer = io.BytesIO()
            output_zip = zipfile.ZipFile(output_buffer, 'w')
            
            updated_xml = document_xml
            rid_counter = max_rid
            
            for i, identifier in enumerate(excel_placeholders):
                if i >= len(excel_file_list):
                    break
                
                excel_filename, excel_base64 = excel_file_list[i]
                placeholder = f'{{{{excel:{identifier}}}}}'
                
                if placeholder in updated_xml:
                    rId = f'rId{rid_counter}'
                    rid_counter += 1
                    
                    ole_xml = self.create_ole_object_xml(rId, excel_filename)
                    updated_xml = updated_xml.replace(placeholder, ole_xml)
                    
                    ole_filename = f'{excel_filename}.bin'
                    rel_entry = f'  <Relationship Id="{rId}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/oleObject" Target="embeddings/{ole_filename}"/>\n'
                    rels_content = rels_content.replace('</Relationships>', rel_entry + '</Relationships>')
                    
                    excel_bytes_data = base64.b64decode(excel_base64)
                    ole_content = self.create_ole_package(excel_bytes_data, excel_filename)
                    output_zip.writestr(f'word/embeddings/{ole_filename}', ole_content)
                    
                    print(f"[OLE] 插入 {excel_filename} 到占位符 {{excel:{identifier}}}")
                else:
                    print(f"[OLE] 警告：未找到占位符 {placeholder}")
            
            for name in template_zip.namelist():
                if name == 'word/document.xml':
                    output_zip.writestr(name, updated_xml.encode('utf-8'))
                elif name == 'word/_rels/document.xml.rels':
                    output_zip.writestr(name, rels_content.encode('utf-8'))
                else:
                    output_zip.writestr(name, template_zip.read(name))
            
            output_zip.close()
            
            print(f"[OLE] 成功处理 {len(excel_file_list)} 个Excel文件")
            return output_buffer.getvalue()
            
        except Exception as e:
            print(f"Pure Python embedding failed: {e}")
            return docx_bytes
    
    def create_ole_package(self, excel_bytes, filename):
        try:
            import olefile
            
            package = io.BytesIO()
            
            with olefile.OleFileIO(package, write_mode=True) as ole:
                ole.create_storage('Root Entry', is_root=True)
                
                internal_stream = io.BytesIO()
                internal_stream.write(b'\x01')
                internal_stream.write(filename.encode('utf-16-le'))
                internal_stream.write(b'\x00\x00')
                internal_stream.write(excel_bytes)
                internal_stream.seek(0)
                
                ole.create_stream('Root Entry', data=internal_stream.getvalue())
            
            package.seek(0)
            return package.getvalue()
            
        except ImportError:
            print("[OLE] olefile not installed")
            return excel_bytes
        except Exception as e:
            print(f"[OLE] olefile error: {e}")
            return excel_bytes
    
    def create_ole_package_fallback(self, excel_bytes, filename):
        package = io.BytesIO()
        
        package.write(b'\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1')
        package.write(b'\x00' * 4)
        
        file_size = len(excel_bytes) + 512
        package.write((file_size).to_bytes(4, 'little'))
        package.write(b'\x00' * 4)
        
        package.write(b'\x00' * 4)
        package.write(b'\x00' * 4)
        
        package.write(b'\x01' + b'\x00' * 3)
        package.write(b'\x00' * 4)
        
        package.write(b'\xFF\xFF\xFF\xFF')
        package.write(b'\x00' * 4)
        
        num_sectors = (len(excel_bytes) + 511) // 512
        package.write((num_sectors).to_bytes(4, 'little'))
        
        package.write(b'\xFF\xFF\xFF\xFF')
        package.write(b'\x00' * 4)
        
        package.write(b'\x00' * 64)
        
        package.write(b'\x00' * 4)
        package.write(b'\x00' * 4)
        
        package.write(b'\x00' * 64)
        
        for i in range(num_sectors):
            start = i * 512
            end = min(start + 512, len(excel_bytes))
            sector_data = excel_bytes[start:end]
            package.write(sector_data)
            if len(sector_data) < 512:
                package.write(b'\x00' * (512 - len(sector_data)))
        
        return package.getvalue()
    
    def create_ole_object_xml(self, relId, filename):
        progid = "Excel.Sheet.12"
        if filename.lower().endswith('.xls'):
            progid = "Excel.Sheet.8"
        
        return f'''<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:v="urn:schemas-microsoft-com:vml" xmlns:o="urn:schemas-microsoft-com:office:office">
  <w:r>
    <w:object w:dxaOrig="4800" w:dyaOrig="2400">
      <v:shape id="_x0000_s1025" type="#_x0000_t75" style="width:400px;height:300px">
        <o:OLEObject Type="Embed" ProgID="{progid}" ShapeID="_x0000_s1025" DrawAspect="Content" ObjectID="_x0000_o1" r:id="{relId}"/>
      </v:shape>
    </w:object>
  </w:r>
</w:p>'''
    
    def process_single_report(self, template_bytes, data_row, community_data):
        excel_files = community_data.get('excel_files', {})
        
        if not excel_files:
            print("[OLE] 没有找到Excel文件")
            return template_bytes
        
        print("[OLE] 使用LibreOffice方式嵌入Excel")
        result = self.embed_excel_libreoffice(template_bytes, excel_files)
        
        if result == template_bytes:
            print("[OLE] LibreOffice方式失败，尝试纯Python方式")
            result = self.embed_excel_pure_python(template_bytes, excel_files, {})
        
        return result
    
    def replace_text_placeholders(self, docx_bytes, data_row):
        try:
            from docx import Document
            import io
            
            doc = Document(io.BytesIO(docx_bytes))
            
            for paragraph in doc.paragraphs:
                for key, value in data_row.items():
                    placeholder = f'{{{{{key}}}}}'
                    if placeholder in paragraph.text:
                        self.replace_in_paragraph(paragraph, placeholder, str(value))
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for key, value in data_row.items():
                                placeholder = f'{{{{{key}}}}}'
                                if placeholder in paragraph.text:
                                    self.replace_in_paragraph(paragraph, placeholder, str(value))
            
            output_buffer = io.BytesIO()
            doc.save(output_buffer)
            return output_buffer.getvalue()
        
        except Exception as e:
            print(f"Error replacing text placeholders: {e}")
            return docx_bytes
    
    def replace_in_paragraph(self, paragraph, placeholder, replacement):
        if placeholder not in paragraph.text:
            return
        
        runs = paragraph.runs
        found = False
        
        for run in runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, replacement)
                found = True
                break
        
        if not found:
            paragraph.text = paragraph.text.replace(placeholder, replacement)
    
    def find_excel_file(self, excel_files, identifier):
        for filename in excel_files:
            if identifier in filename:
                return filename
        return None

def parse_community_zip(zip_bytes):
    result = {
        'excel_files': {},
        'environment_images': {},
        'removed_equipment_images': {}
    }
    
    try:
        with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zf:
            print(f"Total files in zip: {len(zf.namelist())}")
            for name in zf.namelist():
                name = name.replace('\\', '/')
                
                print(f"Checking file: {name}")
                
                if name.endswith('/'):
                    print(f"  -> is directory, skipping")
                    continue
                
                lower_name = name.lower()
                
                parts = lower_name.split('/')
                has_excel_folder = 'excel' in parts or any('excel' in p for p in parts)
                has_env_folder = 'environment_images' in parts or any('environment' in p for p in parts)
                has_equip_folder = 'removedequipment_images' in parts or any('removedequipment' in p for p in parts) or any('removed_equipment' in p for p in parts)
                
                is_excel = has_excel_folder and (name.endswith('.xlsx') or name.endswith('.xls'))
                is_env_image = has_env_folder and (name.endswith('.jpg') or name.endswith('.png') or name.endswith('.jpeg'))
                is_equip_image = has_equip_folder and (name.endswith('.jpg') or name.endswith('.png') or name.endswith('.jpeg'))
                
                if is_excel:
                    filename = os.path.basename(name)
                    with zf.open(name) as f:
                        result['excel_files'][filename] = base64.b64encode(f.read()).decode('utf-8')
                    print(f"  -> ✓ Found excel: {name}")
                
                elif is_env_image:
                    filename = os.path.basename(name)
                    with zf.open(name) as f:
                        result['environment_images'][filename] = base64.b64encode(f.read()).decode('utf-8')
                    print(f"  -> ✓ Found environment image: {name}")
                
                elif is_equip_image:
                    filename = os.path.basename(name)
                    with zf.open(name) as f:
                        result['removed_equipment_images'][filename] = base64.b64encode(f.read()).decode('utf-8')
                    print(f"  -> ✓ Found equipment image: {name}")
                else:
                    print(f"  -> no match (has_excel:{has_excel_folder}, has_env:{has_env_folder}, has_equip:{has_equip_folder})")
    
    except Exception as e:
        print(f"Error parsing zip: {e}")
    
    print(f"Parse result: {len(result['excel_files'])} excel, {len(result['environment_images'])} env images, {len(result['removed_equipment_images'])} equipment images")
    return result

@app.route('/')
def index():
    html_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '退网报告生成小工具.html')
    with open(html_path, 'r', encoding='utf-8') as f:
        content = f.read()
    return content, 200, {'Content-Type': 'text/html; charset=utf-8'}

@app.route('/api/health', methods=['GET'])
def health_check():
    return jsonify({'status': 'ok', 'message': 'Server running'})

@app.route('/api/generate-reports-batch', methods=['POST'])
def generate_reports_batch():
    try:
        data = request.get_json()
        
        template_base64 = data.get('template', '')
        data_rows = data.get('dataRows', [])
        community_zips = data.get('communityZips', [])
        
        if not template_base64:
            return jsonify({'success': False, 'error': '缺少模板文件'}), 400
        
        template_bytes = base64.b64decode(template_base64)
        
        community_data_list = []
        for zip_base64 in community_zips:
            zip_bytes = base64.b64decode(zip_base64)
            community_data = parse_community_zip(zip_bytes)
            community_data_list.append(community_data)
        
        embedder = DocxOLEEmbedder()
        reports = []
        
        for i, row in enumerate(data_rows):
            community_name = row.get('CommunityName', f'小区{i+1}')
            
            community_data = community_data_list[i] if i < len(community_data_list) else (community_data_list[0] if community_data_list else {})
            
            try:
                report_bytes = embedder.process_single_report(template_bytes, row, community_data)
                report_base64 = base64.b64encode(report_bytes).decode('utf-8')
                reports.append({
                    'index': i,
                    'communityName': community_name,
                    'report': report_base64
                })
            except Exception as e:
                print(f"Error processing row {i}: {e}")
                reports.append({
                    'index': i,
                    'communityName': community_name,
                    'error': str(e)
                })
        
        return jsonify({
            'success': True,
            'reports': reports
        })
    
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

if __name__ == '__main__':
    print("Starting server...")
    print("Note: For OLE embedding, Microsoft Office or WPS must be installed")
    port = 8080
    print(f"Server running on http://localhost:{port}")
    try:
        from waitress import serve
        print("Using Waitress production server")
        serve(app, host='0.0.0.0', port=port)
    except ImportError:
        print("Waitress not installed, using Flask development server")
        app.run(debug=False, host='0.0.0.0', port=port)
